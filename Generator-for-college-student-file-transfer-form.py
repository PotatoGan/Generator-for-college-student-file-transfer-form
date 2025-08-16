#!/usr/bin/env python
# -*- coding: utf-8 -*-

import sys
import os
import re
from datetime import datetime
from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QTableWidget, QTableWidgetItem,
                             QFileDialog, QMessageBox, QTabWidget, QLabel, QLineEdit,
                             QGridLayout, QGroupBox, QHeaderView, QAbstractItemView,
                             QDialog, QDialogButtonBox, QFormLayout, QComboBox,
                             QProgressDialog, QTextEdit, QScrollArea)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QTextCursor

class MissingFieldsDialog(QDialog):
    """缺失字段填写对话框"""
    def __init__(self, missing_fields, row_info, parent=None):
        super().__init__(parent)
        self.setWindowTitle("补全缺失字段")
        self.setModal(True)
        self.fields = {}
        
        layout = QVBoxLayout()
        
        # 提示信息
        info_text = f"学号：{row_info.get('学号', 'N/A')}  姓名：{row_info.get('姓名', 'N/A')}  班级：{row_info.get('班级', 'N/A')}"
        info_label = QLabel(info_text)
        info_label.setStyleSheet("font-weight: bold; padding: 10px;")
        layout.addWidget(info_label)
        
        missing_label = QLabel("以下字段在数据中缺失，请补全（可留空）：")
        layout.addWidget(missing_label)
        
        # 字段输入
        form_layout = QFormLayout()
        for field in missing_fields:
            line_edit = QLineEdit()
            
            # 根据字段名称设置提示文本
            if field == '年':
                line_edit.setPlaceholderText('如：2025')
            elif field == '月':
                line_edit.setPlaceholderText('如：7')
            elif field == '日':
                line_edit.setPlaceholderText('如：15')
            elif field == '届':
                line_edit.setPlaceholderText('如：2023')
            elif field == '身份证号':
                line_edit.setPlaceholderText('请输入身份证号')
            elif field == '收档单位名称':
                line_edit.setPlaceholderText('请输入收档单位名称')
            elif field == '转递编号':
                line_edit.setPlaceholderText('请输入转递编号')
            elif field == '生源地':
                line_edit.setPlaceholderText('请输入生源地')
            elif field == '手机号':
                line_edit.setPlaceholderText('请输入手机号')
            elif field == '档案转递类型':
                line_edit.setPlaceholderText('请输入档案转递类型')
                # 添加提示信息
                line_edit.setToolTip('常用类型：转回生源地、签约单位接收、托管单位接收、升学外校接收')
            elif field == '就业单位名称':
                line_edit.setPlaceholderText('请输入就业单位名称')
            elif field == '就业单位地址':
                line_edit.setPlaceholderText('请输入就业单位地址')
            else:
                line_edit.setPlaceholderText(f'请输入{field}')
            
            self.fields[field] = line_edit
            form_layout.addRow(f"{field}:", line_edit)
        
        layout.addLayout(form_layout)
        
        # 添加说明
        note_label = QLabel("提示：点击：“确定”保存填写内容，点击“取消”可选择跳过或留空生成")
        note_label.setStyleSheet("color: gray; font-size: 10pt; padding: 5px;")
        layout.addWidget(note_label)
        
        # 按钮
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | 
            QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        
        self.setLayout(layout)
        self.resize(450, min(350 + len(missing_fields) * 35, 650))
    
    def get_values(self):
        """获取用户输入的值"""
        values = {}
        for field, edit in self.fields.items():
            value = edit.text().strip()
            values[field] = value
        return values

class WordGeneratorThread(QThread):
    """Word文档生成线程"""
    progress = pyqtSignal(int)
    status = pyqtSignal(str)
    finished = pyqtSignal(int)
    error = pyqtSignal(str)
    
    def __init__(self, data_rows, template_path, output_dir):
        super().__init__()
        self.data_rows = data_rows
        self.template_path = template_path
        self.output_dir = output_dir
        self.success_count = 0
        
    def run(self):
        try:
            total = len(self.data_rows)
            for i, row_data in enumerate(self.data_rows):
                self.status.emit(f"正在生成：{row_data.get('姓名', 'unknown')}")
                
                # 生成文档
                self.generate_single_doc(row_data)
                self.success_count += 1
                
                # 更新进度
                progress = int((i + 1) / total * 100)
                self.progress.emit(progress)
                
            self.finished.emit(self.success_count)
        except Exception as e:
            self.error.emit(str(e))
    
    def generate_single_doc(self, data):
        """生成单个文档"""
        doc = Document(self.template_path)
        
        # 替换文档中的所有占位符
        for paragraph in doc.paragraphs:
            self.replace_text_in_paragraph(paragraph, data)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_paragraph(paragraph, data)
        
        # 生成文件名
        filename = f"{data.get('学号', 'unknown')}_{data.get('姓名', 'unknown')}_{data.get('班级', 'unknown')}.docx"
        # 清理文件名中的非法字符
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        
        # 保存文档
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
    
    def replace_text_in_paragraph(self, paragraph, data):
        """替换段落中的占位符，保持原有格式"""
        # 先检查整个段落文本中是否包含占位符
        full_text = ''.join(run.text for run in paragraph.runs)
        
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                # 占位符可能被分割在多个run中，需要特殊处理
                self.replace_placeholder_in_runs(paragraph.runs, placeholder, str(value) if value else '')
    
    def replace_placeholder_in_runs(self, runs, placeholder, replacement):
        """在runs中替换占位符，处理占位符被分割的情况"""
        text = ''
        run_info = []
        
        # 收集所有run的文本和信息
        for run in runs:
            run_start = len(text)
            text += run.text
            run_end = len(text)
            run_info.append({
                'run': run,
                'start': run_start,
                'end': run_end,
                'original_text': run.text
            })
        
        # 查找并替换占位符
        new_text = text.replace(placeholder, replacement)
        
        # 如果没有变化，直接返回
        if new_text == text:
            return
        
        # 计算需要调整的位置
        placeholder_pos = text.find(placeholder)
        if placeholder_pos == -1:
            return
        
        # 找到占位符涉及的runs
        affected_runs = []
        for info in run_info:
            if info['start'] <= placeholder_pos < info['end'] or \
               info['start'] < placeholder_pos + len(placeholder) <= info['end'] or \
               (info['start'] >= placeholder_pos and info['end'] <= placeholder_pos + len(placeholder)):
                affected_runs.append(info)
        
        if affected_runs:
            # 在第一个受影响的run中进行替换
            first_run = affected_runs[0]['run']
            
            # 构建新的文本
            before_placeholder = text[:placeholder_pos]
            after_placeholder = text[placeholder_pos + len(placeholder):]
            
            # 计算第一个run应该包含的文本
            first_run_start = affected_runs[0]['start']
            first_run_text_before = before_placeholder[first_run_start:] if first_run_start < len(before_placeholder) else ''
            
            # 设置第一个run的新文本
            first_run.text = first_run_text_before + replacement
            
            # 处理剩余的文本
            remaining_text_start = placeholder_pos + len(placeholder)
            if len(affected_runs) > 1:
                last_run = affected_runs[-1]
                if last_run['end'] > remaining_text_start:
                    # 如果最后一个受影响的run还有剩余文本
                    remaining_in_last = after_placeholder[:last_run['end'] - remaining_text_start]
                    first_run.text += remaining_in_last
                
                # 清空中间的runs
                for info in affected_runs[1:]:
                    info['run'].text = ''

class DocumentGenerator:
    """文档生成器类 - 提取通用的文档处理逻辑"""
    
    @staticmethod
    def replace_text_in_paragraph(paragraph, data):
        """替换段落中的占位符，保持原有格式"""
        # 先检查整个段落文本中是否包含占位符
        full_text = ''.join(run.text for run in paragraph.runs)
        
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in full_text:
                # 占位符可能被分割在多个run中，需要特殊处理
                DocumentGenerator.replace_placeholder_in_runs(paragraph.runs, placeholder, str(value) if value else '')
    
    @staticmethod
    def replace_placeholder_in_runs(runs, placeholder, replacement):
        """在runs中替换占位符，处理占位符被分割的情况"""
        text = ''
        run_info = []
        
        # 收集所有run的文本和信息
        for run in runs:
            run_start = len(text)
            text += run.text
            run_end = len(text)
            run_info.append({
                'run': run,
                'start': run_start,
                'end': run_end,
                'original_text': run.text
            })
        
        # 查找并替换占位符
        new_text = text.replace(placeholder, replacement)
        
        # 如果没有变化，直接返回
        if new_text == text:
            return
        
        # 计算需要调整的位置
        placeholder_pos = text.find(placeholder)
        if placeholder_pos == -1:
            return
        
        # 找到占位符涉及的runs
        affected_runs = []
        for info in run_info:
            if info['start'] <= placeholder_pos < info['end'] or \
               info['start'] < placeholder_pos + len(placeholder) <= info['end'] or \
               (info['start'] >= placeholder_pos and info['end'] <= placeholder_pos + len(placeholder)):
                affected_runs.append(info)
        
        if affected_runs:
            # 在第一个受影响的run中进行替换
            first_run = affected_runs[0]['run']
            
            # 构建新的文本
            before_placeholder = text[:placeholder_pos]
            after_placeholder = text[placeholder_pos + len(placeholder):]
            
            # 计算第一个run应该包含的文本
            first_run_start = affected_runs[0]['start']
            first_run_text_before = before_placeholder[first_run_start:] if first_run_start < len(before_placeholder) else ''
            
            # 设置第一个run的新文本
            first_run.text = first_run_text_before + replacement
            
            # 处理剩余的文本
            remaining_text_start = placeholder_pos + len(placeholder)
            if len(affected_runs) > 1:
                last_run = affected_runs[-1]
                if last_run['end'] > remaining_text_start:
                    # 如果最后一个受影响的run还有剩余文本
                    remaining_in_last = after_placeholder[:last_run['end'] - remaining_text_start]
                    first_run.text += remaining_in_last
                
                # 清空中间的runs
                for info in affected_runs[1:]:
                    info['run'].text = ''
    
    @staticmethod
    def generate_document(template_path, data, output_path):
        """生成单个文档的通用方法"""
        try:
            # 加载模板
            doc = Document(template_path)
            
            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                DocumentGenerator.replace_text_in_paragraph(paragraph, data)
            
            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            DocumentGenerator.replace_text_in_paragraph(paragraph, data)
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            raise Exception(f"生成文档时出错：{str(e)}")

class ArchiveTransferGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        self.excel_data = None
        self.template_variables = None
        self.initUI()
        
    def initUI(self):
        self.setWindowTitle('档案转递文档批量生成工具')
        self.setGeometry(100, 100, 1200, 700)
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QVBoxLayout()
        central_widget.setLayout(main_layout)
        
        # 创建选项卡
        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)
        
        # Excel导入页面
        self.excel_tab = self.create_excel_tab()
        self.tab_widget.addTab(self.excel_tab, "Excel批量生成")
        
        # 手动填写页面
        self.manual_tab = self.create_manual_tab()
        self.tab_widget.addTab(self.manual_tab, "手动填写生成")
        
        # 关于页面
        self.about_tab = self.create_about_tab()
        self.tab_widget.addTab(self.about_tab, "关于")
        
        # 状态栏
        self.statusBar().showMessage('就绪')
        
    def create_excel_tab(self):
        """创建Excel导入页面"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # 按钮组
        button_layout = QHBoxLayout()
        
        self.load_excel_btn = QPushButton('📁 读取Excel文件')
        self.load_excel_btn.clicked.connect(self.load_excel)
        button_layout.addWidget(self.load_excel_btn)
        
        self.select_all_btn = QPushButton('☑ 全选')
        self.select_all_btn.clicked.connect(self.select_all)
        self.select_all_btn.setEnabled(False)
        button_layout.addWidget(self.select_all_btn)
        
        self.deselect_all_btn = QPushButton('☐ 取消全选')
        self.deselect_all_btn.clicked.connect(self.deselect_all)
        self.deselect_all_btn.setEnabled(False)
        button_layout.addWidget(self.deselect_all_btn)
        
        self.generate_btn = QPushButton('📄 批量生成Word文档')
        self.generate_btn.clicked.connect(self.batch_generate)
        self.generate_btn.setEnabled(False)
        button_layout.addWidget(self.generate_btn)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        # 数据表格
        self.data_table = QTableWidget()
        self.data_table.setAlternatingRowColors(True)
        self.data_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.data_table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked)
        # 当单元格内容改变时更新转档字号
        self.data_table.itemChanged.connect(self.on_table_item_changed)
        layout.addWidget(self.data_table)
        
        tab.setLayout(layout)
        return tab
    
    def create_manual_tab(self):
        """创建手动填写页面"""
        tab = QWidget()
        layout = QVBoxLayout()
        
        # 创建滚动区域以容纳所有字段
        form_group = QGroupBox("档案信息填写")
        form_layout = QGridLayout()
        
        # 定义所有可能的字段
        self.manual_fields = {}
        field_list = [
            ('姓名', '请输入姓名'),
            ('学号', '请输入学号'),
            ('班级', '请输入班级'),
            ('届', '如：2023'),
            ('年', '如：2025'),
            ('月', '如：7'),
            ('日', '如：15'),
            ('身份证号', '请输入身份证号'),
            ('收档单位名称', '请输入收档单位名称'),
            ('转递编号', '请输入转递编号'),
            ('生源地', '请输入生源地'),
            ('手机号', '请输入手机号'),
            ('档案转递类型', None),  # 将使用下拉框
            ('就业单位名称', '请输入就业单位名称'),
            ('就业单位地址', '请输入就业单位地址'),
        ]
        
        row = 0
        col = 0
        for field_name, placeholder in field_list:
            label = QLabel(f"{field_name}:")
            
            if field_name == '档案转递类型':
                # 创建下拉框
                combo_box = QComboBox()
                combo_box.addItems([
                    '',  # 空选项
                    '转回生源地',
                    '签约单位接收',
                    '托管单位接收',
                    '升学外校接收'
                ])
                combo_box.setEditable(True)  # 允许自定义输入
                combo_box.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)  # 不自动添加到列表
                self.manual_fields[field_name] = combo_box
                form_layout.addWidget(label, row, col * 2)
                form_layout.addWidget(combo_box, row, col * 2 + 1)
            else:
                # 创建普通输入框
                line_edit = QLineEdit()
                line_edit.setPlaceholderText(placeholder)
                self.manual_fields[field_name] = line_edit
                form_layout.addWidget(label, row, col * 2)
                form_layout.addWidget(line_edit, row, col * 2 + 1)
            
            col += 1
            if col >= 3:  # 每行3个字段
                col = 0
                row += 1
        
        form_group.setLayout(form_layout)
        layout.addWidget(form_group)
        
        # 按钮
        button_layout = QHBoxLayout()
        
        self.clear_btn = QPushButton('🔄 清空所有字段')
        self.clear_btn.clicked.connect(self.clear_manual_fields)
        button_layout.addWidget(self.clear_btn)
        
        self.generate_single_btn = QPushButton('📄 生成Word文档')
        self.generate_single_btn.clicked.connect(self.generate_single)
        button_layout.addWidget(self.generate_single_btn)
        
        # 快速填充今天日期
        self.fill_today_btn = QPushButton('📅 填充今天日期')
        self.fill_today_btn.clicked.connect(self.fill_today_date)
        button_layout.addWidget(self.fill_today_btn)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        tab.setLayout(layout)
        return tab
    
    def create_about_tab(self):
        """创建关于页面"""
        tab = QWidget()
        main_layout = QVBoxLayout()
        
        # 创建滚动区域
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        
        # 内容容器
        content_widget = QWidget()
        content_layout = QVBoxLayout()
        
        # 标题
        title = QLabel("档案转递文档批量生成工具")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size: 18pt; font-weight: bold; padding: 10px;")
        content_layout.addWidget(title)
        
        # 作者信息组
        author_group = QGroupBox("作者信息")
        author_layout = QVBoxLayout()
        
        author_info = """
        <p><b>作者：</b>周天僖</p>
        <p><b>联系邮箱：</b>2023520354@bipt.edu.cn</p>
        <p><b>制作时间：</b>2025年8月</p>
        <p><b>版本：</b>v1.0</p>
        """
        author_label = QLabel(author_info)
        author_label.setTextFormat(Qt.TextFormat.RichText)
        author_label.setStyleSheet("padding: 10px;")
        author_layout.addWidget(author_label)
        author_group.setLayout(author_layout)
        content_layout.addWidget(author_group)
        
        # 使用说明组
        manual_group = QGroupBox("软件使用说明")
        manual_layout = QVBoxLayout()
        
        manual_text = QTextEdit()
        manual_text.setReadOnly(True)
        manual_text.setHtml("""
        <h3>一、准备工作</h3>
        <h4>1. 文件夹结构</h4>
        <p>请确保程序所在目录下有以下文件夹结构：</p>
        <ul>
            <li><b>template/</b> - 存放Word模板文件（.docx格式）</li>
            <li>模板文件中使用 <b>{{字段名}}</b> 作为占位符</li>
        </ul>
        
        <h4>2. Word模板格式</h4>
        <p>模板中的占位符格式必须为：<b>{{字段名}}</b></p>
        <p>常用字段包括：</p>
        <ul>
            <li>{{姓名}} - 学生姓名</li>
            <li>{{学号}} - 学生学号</li>
            <li>{{班级}} - 班级名称</li>
            <li>{{届}} - 毕业届别（如2023）</li>
            <li>{{年}}、{{月}}、{{日}} - 日期信息</li>
            <li>{{转档字号}} - 自动生成的转档编号</li>
            <li>{{身份证号}} - 身份证号码</li>
            <li>{{收档单位名称}} - 接收档案的单位</li>
            <li>{{手机号}} - 联系电话</li>
            <li>其他自定义字段...</li>
        </ul>
        
        <h4>3. Excel数据文件</h4>
        <p>Excel文件应包含与模板对应的列名，程序会自动匹配。</p>
        <p>如果Excel中有日期相关字段（如"提交时间"、"提交日期"等），程序会自动提取年、月、日。</p>
        
        <hr>
        
        <h3>二、Excel批量生成功能</h3>
        
        <h4>1. 导入Excel数据</h4>
        <ol>
            <li>点击 <b>"📁 读取Excel文件"</b> 按钮</li>
            <li>选择包含学生信息的Excel文件（.xlsx或.xls格式）</li>
            <li>程序会自动读取数据并显示在表格中</li>
            <li>如果有日期字段，会自动提取年、月、日信息</li>
            <li>自动生成转档字号（格式：年份后两位+学号_班级）</li>
        </ol>
        
        <h4>2. 编辑数据</h4>
        <ul>
            <li><b>双击单元格</b>可以直接编辑内容</li>
            <li>修改年、学号或班级后，转档字号会自动更新</li>
            <li>所有修改都会实时反映在生成的文档中</li>
        </ul>
        
        <h4>3. 选择要生成的记录</h4>
        <ul>
            <li>每行前有复选框，勾选要生成文档的记录</li>
            <li>使用 <b>"☑ 全选"</b> 选择所有记录</li>
            <li>使用 <b>"☐ 取消全选"</b> 取消所有选择</li>
        </ul>
        
        <h4>4. 批量生成文档</h4>
        <ol>
            <li>选择要生成文档的记录（打勾）</li>
            <li>点击 <b>"📄 批量生成Word文档"</b> 按钮</li>
            <li>选择输出文件夹</li>
            <li>如果有缺失字段，会逐条弹出对话框：
                <ul>
                    <li>可以填写缺失的信息</li>
                    <li>可以留空继续生成</li>
                    <li>可以选择跳过该记录</li>
                </ul>
            </li>
            <li>等待生成完成，进度条会显示当前进度</li>
        </ol>
        
        <h4>5. 文件命名规则</h4>
        <p>生成的文件名格式：<b>学号_姓名_班级.docx</b></p>
        
        <hr>
        
        <h3>三、手动填写生成功能</h3>
        
        <h4>1. 填写信息</h4>
        <ol>
            <li>切换到 <b>"手动填写生成"</b> 标签页</li>
            <li>在相应的输入框中填写信息</li>
            <li>必填字段：姓名、学号、班级</li>
            <li>其他字段可选填</li>
        </ol>
        
        <h4>2. 快速功能</h4>
        <ul>
            <li><b>"📅 填充今天日期"</b> - 自动填入当前日期的年、月、日</li>
            <li><b>"🔄 清空所有字段"</b> - 清空所有已填写的内容</li>
        </ul>
        
        <h4>3. 生成单个文档</h4>
        <ol>
            <li>填写必要信息</li>
            <li>点击 <b>"📄 生成Word文档"</b> 按钮</li>
            <li>选择输出文件夹</li>
            <li>如果有未填写的字段，会提示是否继续（未填写的字段在文档中留空）</li>
        </ol>
        
        <hr>
        
        <h3>四、特殊字段说明</h3>
        
        <h4>1. 年份处理</h4>
        <ul>
            <li>在文档中显示完整年份（如：2025）</li>
            <li>转档字号只使用年份后两位（如：25）</li>
        </ul>
        
        <h4>2. 转档字号</h4>
        <ul>
            <li>自动生成，格式：年份后两位+学号_班级</li>
            <li>例如：25202201_计算机1班</li>
            <li>修改年、学号或班级后自动更新</li>
        </ul>
        
        <h4>3. 日期提取</h4>
        <p>程序会自动识别以下格式的日期字段：</p>
        <ul>
            <li>提交时间、提交日期</li>
            <li>日期、时间</li>
            <li>创建时间、更新时间</li>
        </ul>
        <p>支持的日期格式：</p>
        <ul>
            <li>2025/8/15 或 2025-08-15</li>
            <li>包含时间的格式也可识别</li>
        </ul>
        
        <hr>
        
        <h3>五、常见问题</h3>
        
        <h4>Q: 为什么提示找不到template文件夹？</h4>
        <p>A: 请在程序所在目录创建名为"template"的文件夹，并放入Word模板文件。</p>
        
        <h4>Q: 生成的文档中有{{xxx}}这样的文字？</h4>
        <p>A: 这表示该字段在数据中缺失，请检查Excel数据或在弹出的对话框中补全。</p>
        
        <h4>Q: 如何修改已导入的数据？</h4>
        <p>A: 直接双击表格中的单元格即可编辑，修改后的内容会直接用于生成文档。</p>
        
        <h4>Q: 转档字号格式可以自定义吗？</h4>
        <p>A: 目前转档字号格式固定为：年份后两位+学号_班级。如需其他格式，可手动编辑表格中的转档字号列。</p>
        
        <h4>Q: 批量生成时某些记录失败了怎么办？</h4>
        <p>A: 程序会显示成功生成的数量，失败的记录可以单独检查并重新生成。</p>
        
        <hr>
        
        <h3>六、使用技巧</h3>
        
        <ol>
            <li><b>批量处理前先检查：</b>导入Excel后，先浏览一遍数据，确认关键字段都有值。</li>
            <li><b>利用全选功能：</b>如果大部分记录都要生成，先全选再取消个别不需要的。</li>
            <li><b>表格编辑即时生效：</b>发现错误可以直接在表格中修改，不需要重新导入。</li>
            <li><b>缺失字段灵活处理：</b>不是所有字段都必须填写，可以根据实际需要选择补全或留空。</li>
            <li><b>保存编辑结果：</b>如果对表格做了大量修改，建议生成文档后也导出一份修改后的Excel作为备份。</li>
        </ol>
        
        <hr>
        
        <p style="text-align: center; color: gray; margin-top: 20px;">
        如有问题或建议，请联系：2023520354@bipt.edu.cn
        </p>
        """)
        
        manual_layout.addWidget(manual_text)
        manual_group.setLayout(manual_layout)
        content_layout.addWidget(manual_group)
        
        # 设置内容widget
        content_widget.setLayout(content_layout)
        scroll_area.setWidget(content_widget)
        
        # 添加到主布局
        main_layout.addWidget(scroll_area)
        tab.setLayout(main_layout)
        return tab
    
    def on_table_item_changed(self, item):
        """当表格项改变时触发"""
        if not item:
            return
        
        row = item.row()
        col = item.column()
        
        # 获取列标题
        header = self.data_table.horizontalHeaderItem(col)
        if not header:
            return
        
        column_name = header.text()
        
        # 如果修改的是年、学号或班级列，更新转档字号
        if column_name in ['年', '学号', '班级']:
            self.update_transfer_number_for_row(row)
    
    def update_transfer_number_for_row(self, row):
        """更新指定行的转档字号"""
        # 查找年、学号、班级列的索引
        year_col = month_col = day_col = student_id_col = class_col = transfer_col = -1
        
        for col in range(self.data_table.columnCount()):
            header = self.data_table.horizontalHeaderItem(col)
            if header:
                header_text = header.text()
                if header_text == '年':
                    year_col = col
                elif header_text == '月':
                    month_col = col
                elif header_text == '日':
                    day_col = col
                elif header_text == '学号':
                    student_id_col = col
                elif header_text == '班级':
                    class_col = col
                elif header_text == '转档字号':
                    transfer_col = col
        
        # 如果找到了所有必要的列
        if year_col >= 0 and student_id_col >= 0 and class_col >= 0:
            year_item = self.data_table.item(row, year_col)
            student_id_item = self.data_table.item(row, student_id_col)
            class_item = self.data_table.item(row, class_col)
            
            if year_item and student_id_item and class_item:
                year = year_item.text().strip()
                student_id = student_id_item.text().strip()
                class_name = class_item.text().strip()
                
                if year and student_id and class_name:
                    # 生成转档字号：年份后两位 + 学号 + _ + 班级
                    year_suffix = year[-2:] if len(year) >= 2 else year
                    transfer_number = f"{year_suffix}{student_id}_{class_name}"
                    
                    # 如果转档字号列存在，更新它
                    if transfer_col >= 0:
                        transfer_item = QTableWidgetItem(transfer_number)
                        self.data_table.setItem(row, transfer_col, transfer_item)
    
    def load_excel(self):
        """加载Excel文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, 
            "选择Excel文件", 
            "", 
            "Excel Files (*.xlsx *.xls)"
        )
        
        if file_path:
            try:
                # 读取Excel文件
                self.excel_data = pd.read_excel(file_path)
                self.excel_data = self.excel_data.fillna('')  # 将NaN替换为空字符串
                
                # 处理日期字段，提取年月日
                self.process_date_fields()
                
                # 显示数据到表格
                self.display_data()
                
                # 启用按钮
                self.select_all_btn.setEnabled(True)
                self.deselect_all_btn.setEnabled(True)
                self.generate_btn.setEnabled(True)
                
                self.statusBar().showMessage(f'已加载 {len(self.excel_data)} 条记录')
                
            except Exception as e:
                QMessageBox.critical(self, "错误", f"读取Excel文件失败：\n{str(e)}")
    
    def process_date_fields(self):
        """处理日期字段，提取年月日"""
        if self.excel_data is None:
            return
        
        # 查找日期字段
        date_field_names = ['提交时间', '提交日期', '日期', '时间', '创建时间', '更新时间']
        date_field = None
        
        for field_name in date_field_names:
            if field_name in self.excel_data.columns:
                date_field = field_name
                break
        
        if not date_field:
            return
        
        # 提取年月日
        for idx, row in self.excel_data.iterrows():
            date_value = row[date_field]
            if date_value and not pd.isna(date_value):
                try:
                    # 如果是datetime对象
                    if hasattr(date_value, 'year'):
                        self.excel_data.at[idx, '年'] = str(date_value.year)  # 完整年份
                        self.excel_data.at[idx, '月'] = str(date_value.month)
                        self.excel_data.at[idx, '日'] = str(date_value.day)
                    else:
                        # 如果是字符串
                        date_str = str(date_value).strip()
                        if ' ' in date_str:
                            date_part = date_str.split(' ')[0]
                        else:
                            date_part = date_str
                        
                        if '/' in date_part:
                            parts = date_part.split('/')
                            if len(parts) >= 3:
                                self.excel_data.at[idx, '年'] = parts[0].strip()  # 完整年份
                                self.excel_data.at[idx, '月'] = str(int(parts[1])) if parts[1].isdigit() else parts[1]
                                self.excel_data.at[idx, '日'] = str(int(parts[2])) if parts[2].isdigit() else parts[2]
                        elif '-' in date_part:
                            parts = date_part.split('-')
                            if len(parts) >= 3:
                                self.excel_data.at[idx, '年'] = parts[0].strip()  # 完整年份
                                self.excel_data.at[idx, '月'] = str(int(parts[1])) if parts[1].isdigit() else parts[1]
                                self.excel_data.at[idx, '日'] = str(int(parts[2])) if parts[2].isdigit() else parts[2]
                except:
                    pass
        
        # 生成转档字号
        if '年' in self.excel_data.columns and '学号' in self.excel_data.columns and '班级' in self.excel_data.columns:
            for idx, row in self.excel_data.iterrows():
                year = str(row['年']) if pd.notna(row['年']) else ''
                student_id = str(row['学号']) if pd.notna(row['学号']) else ''
                class_name = str(row['班级']) if pd.notna(row['班级']) else ''
                
                if year and student_id and class_name:
                    # 转档字号使用年份后两位
                    year_suffix = year[-2:] if len(year) >= 2 else year
                    self.excel_data.at[idx, '转档字号'] = f"{year_suffix}{student_id}_{class_name}"
    
    def display_data(self):
        """显示数据到表格"""
        if self.excel_data is None:
            return
        
        # 设置表格
        self.data_table.setRowCount(len(self.excel_data))
        self.data_table.setColumnCount(len(self.excel_data.columns) + 1)
        
        # 设置表头
        headers = ['选择'] + list(self.excel_data.columns)
        self.data_table.setHorizontalHeaderLabels(headers)
        
        # 填充数据
        for row_idx, row_data in self.excel_data.iterrows():
            # 添加复选框
            checkbox = QTableWidgetItem()
            checkbox.setCheckState(Qt.CheckState.Unchecked)
            self.data_table.setItem(row_idx, 0, checkbox)
            
            # 添加数据
            for col_idx, value in enumerate(row_data):
                # 处理各种数据类型
                if pd.isna(value):
                    item_text = ''
                elif isinstance(value, (pd.Timestamp, datetime)):
                    # 格式化日期时间显示
                    item_text = value.strftime('%Y/%m/%d %H:%M:%S') if hasattr(value, 'strftime') else str(value)
                else:
                    item_text = str(value)
                item = QTableWidgetItem(item_text)
                self.data_table.setItem(row_idx, col_idx + 1, item)
        
        # 调整列宽
        self.data_table.resizeColumnsToContents()
        self.data_table.horizontalHeader().setStretchLastSection(True)
    
    def select_all(self):
        """全选"""
        for row in range(self.data_table.rowCount()):
            item = self.data_table.item(row, 0)
            if item:
                item.setCheckState(Qt.CheckState.Checked)
    
    def deselect_all(self):
        """取消全选"""
        for row in range(self.data_table.rowCount()):
            item = self.data_table.item(row, 0)
            if item:
                item.setCheckState(Qt.CheckState.Unchecked)
    
    def get_template_variables(self):
        """获取模板中的变量"""
        template_dir = Path('./template')
        if not template_dir.exists():
            template_dir = Path('./模板')
            if not template_dir.exists():
                QMessageBox.warning(self, "警告", "未找到template文件夹，请确保模板文件夹存在")
                return None
        
        # 查找Word模板文件
        template_files = list(template_dir.glob('*.docx'))
        if not template_files:
            QMessageBox.warning(self, "警告", "template文件夹中未找到Word模板文件")
            return None
        
        template_path = template_files[0]
        
        try:
            doc = Document(template_path)
            variables = set()
            
            # 查找段落中的变量
            for paragraph in doc.paragraphs:
                # 获取段落的完整文本（合并所有runs）
                full_text = ''.join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
                variables.update(re.findall(r'\{\{(\w+)\}\}', full_text))
            
            # 查找表格中的变量
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # 获取段落的完整文本（合并所有runs）
                            full_text = ''.join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
                            variables.update(re.findall(r'\{\{(\w+)\}\}', full_text))
            
            return template_path, variables
        except Exception as e:
            QMessageBox.critical(self, "错误", f"读取模板文件失败：\n{str(e)}")
            return None
    
    def get_row_data_from_table(self, row_idx):
        """从表格获取指定行的数据"""
        data = {}
        
        # 获取所有列的数据（跳过第一列的复选框）
        for col in range(1, self.data_table.columnCount()):
            header = self.data_table.horizontalHeaderItem(col)
            if header:
                column_name = header.text()
                item = self.data_table.item(row_idx, col)
                if item:
                    data[column_name] = item.text()
                else:
                    data[column_name] = ''
        
        return data
    
    def batch_generate(self):
        """批量生成Word文档"""
        # 获取选中的行
        selected_rows = []
        for row in range(self.data_table.rowCount()):
            if self.data_table.item(row, 0).checkState() == Qt.CheckState.Checked:
                selected_rows.append(row)
        
        if not selected_rows:
            QMessageBox.warning(self, "警告", "请至少选择一行数据")
            return
        
        # 获取模板变量
        template_info = self.get_template_variables()
        if not template_info:
            return
        
        template_path, template_variables = template_info
        
        # 选择输出目录
        output_dir = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if not output_dir:
            return
        
        # 准备数据 - 直接从表格获取数据
        data_rows = []
        
        for row_idx in selected_rows:
            # 直接从表格获取当前显示的数据
            row_data = self.get_row_data_from_table(row_idx)
            
            # 检查缺失的必要字段
            missing_fields = set()
            for field in template_variables:
                if field not in row_data or not row_data[field]:
                    missing_fields.add(field)
            
            # 如果有缺失字段，弹出对话框让用户补全
            if missing_fields:
                # 显示当前记录信息
                info_for_dialog = {
                    '姓名': row_data.get('姓名', 'N/A'),
                    '学号': row_data.get('学号', 'N/A'),
                    '班级': row_data.get('班级', 'N/A')
                }
                
                dialog = MissingFieldsDialog(missing_fields, info_for_dialog, self)
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    # 获取用户填写的值
                    filled_values = dialog.get_values()
                    
                    # 更新数据
                    for field, value in filled_values.items():
                        if value:  # 只更新非空值
                            row_data[field] = value
                            
                            # 同时更新表格显示
                            for col in range(1, self.data_table.columnCount()):
                                header = self.data_table.horizontalHeaderItem(col)
                                if header and header.text() == field:
                                    item = QTableWidgetItem(value)
                                    self.data_table.setItem(row_idx, col, item)
                                    break
                    
                    # 如果用户填写了年、学号或班级，更新转档字号
                    if any(key in filled_values for key in ['年', '学号', '班级']):
                        self.update_transfer_number_for_row(row_idx)
                        # 重新获取更新后的数据
                        row_data = self.get_row_data_from_table(row_idx)
                else:
                    # 用户取消了，但仍然可以选择继续（字段留空）
                    reply = QMessageBox.question(
                        self,
                        "跳过此记录",
                        f"学号：{info_for_dialog['学号']} 姓名：{info_for_dialog['姓名']}\n\n是否跳过此记录？\n\n选择“是”跳过此记录，选择“否”将缺失字段留空继续生成。",
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                    )
                    if reply == QMessageBox.StandardButton.Yes:
                        continue  # 跳过这条记录
                    # 否则继续，缺失字段留空
            
            data_rows.append(row_data)
        
        if not data_rows:
            QMessageBox.warning(self, "警告", "没有要生成的数据")
            return
        
        # 创建进度对话框
        progress_dialog = QProgressDialog("正在生成Word文档...", "取消", 0, 100, self)
        progress_dialog.setWindowTitle("批量生成进度")
        progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
        progress_dialog.show()
        
        # 创建并启动生成线程
        self.generator_thread = WordGeneratorThread(data_rows, template_path, output_dir)
        self.generator_thread.progress.connect(progress_dialog.setValue)
        self.generator_thread.status.connect(lambda msg: progress_dialog.setLabelText(msg))
        self.generator_thread.finished.connect(lambda count: self.on_generation_finished(count, progress_dialog))
        self.generator_thread.error.connect(lambda msg: self.on_generation_error(msg, progress_dialog))
        self.generator_thread.start()
    
    def on_generation_finished(self, count, progress_dialog):
        """生成完成处理"""
        progress_dialog.close()
        QMessageBox.information(self, "完成", f"成功生成 {count} 个Word文档")
        self.statusBar().showMessage(f'成功生成 {count} 个文档')
    
    def on_generation_error(self, error_msg, progress_dialog):
        """生成错误处理"""
        progress_dialog.close()
        QMessageBox.critical(self, "错误", f"生成文档时出错：\n{error_msg}")
        self.statusBar().showMessage('生成失败')
    
    def clear_manual_fields(self):
        """清空手动填写的字段"""
        for field_name, field_widget in self.manual_fields.items():
            if isinstance(field_widget, QComboBox):
                field_widget.setCurrentIndex(0)  # 重置到第一个空选项
            else:
                field_widget.clear()
    
    def fill_today_date(self):
        """填充今天的日期"""
        today = datetime.now()
        if '年' in self.manual_fields:
            self.manual_fields['年'].setText(str(today.year))  # 完整年份
        if '月' in self.manual_fields:
            self.manual_fields['月'].setText(str(today.month))
        if '日' in self.manual_fields:
            self.manual_fields['日'].setText(str(today.day))
    
    def generate_single(self):
        """生成单个Word文档 - 修复版本"""
        # 获取模板信息
        template_info = self.get_template_variables()
        if not template_info:
            return
        
        template_path, template_variables = template_info
        
        # 收集填写的数据
        data = {}
        for field_name, field_widget in self.manual_fields.items():
            if isinstance(field_widget, QComboBox):
                value = field_widget.currentText().strip()
            else:
                value = field_widget.text().strip()
            
            if value:
                data[field_name] = value
        
        # 检查必填字段
        required_fields = ['姓名', '学号', '班级']
        missing_required = [f for f in required_fields if not data.get(f)]
        if missing_required:
            QMessageBox.warning(self, "警告", f"请填写必填字段：{', '.join(missing_required)}")
            return
        
        # 生成转档字号（使用年份后两位）
        if '年' in data and '学号' in data and '班级' in data:
            year = data['年']
            year_suffix = year[-2:] if len(year) >= 2 else year
            data['转档字号'] = f"{year_suffix}{data['学号']}_{data['班级']}"
        
        # 检查模板中的其他变量
        missing_fields = template_variables - set(data.keys())
        if missing_fields:
            reply = QMessageBox.question(
                self, 
                "缺失字段", 
                f"以下字段未填写：{', '.join(missing_fields)}\n是否继续生成？",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                return
            
            # 将缺失字段设为空
            for field in missing_fields:
                data[field] = ''
        
        # 选择输出目录
        output_dir = QFileDialog.getExistingDirectory(self, "选择输出目录")
        if not output_dir:
            return
        
        try:
            # 生成文件名
            filename = f"{data.get('学号', 'unknown')}_{data.get('姓名', 'unknown')}_{data.get('班级', 'unknown')}.docx"
            filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            output_path = os.path.join(output_dir, filename)
            
            # 使用统一的文档生成方法
            DocumentGenerator.generate_document(template_path, data, output_path)
            
            QMessageBox.information(self, "成功", f"文档已生成：\n{filename}")
            self.statusBar().showMessage(f'文档已生成：{filename}')
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成文档失败：\n{str(e)}")

def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    
    # 设置应用程序字体
    font = QFont()
    font.setPointSize(10)
    app.setFont(font)
    
    window = ArchiveTransferGenerator()
    window.show()
    
    sys.exit(app.exec())

if __name__ == '__main__':
    main()